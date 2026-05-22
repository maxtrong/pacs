# -*- coding: utf-8 -*-
import os
import cv2
import numpy as np
import pydicom
import customtkinter as ctk
from PIL import Image, ImageTk, ImageDraw
from tkinter import filedialog, messagebox, scrolledtext
from datetime import datetime
import json
import shutil
import time

# Configuración de UI
ctk.set_appearance_mode("dark")
CORNER_RADIUS = 12 
FONT_NORMAL = ("Segoe UI", 13)
FONT_BOLD = ("Segoe UI", 13, "bold")
FONT_ICONS = ("Segoe UI", 24)

class EstacionRadiologica(ctk.CTk):
    def __init__(self):
        super().__init__()
        # Ocultar la ventana principal mientras está el Splash
        self.withdraw()
        
        # --- LANZAR SPLASH SCREEN ---
        self.mostrar_splash()

    def mostrar_splash(self):
        # Crear ventana de bienvenida (Splash)
        self.splash = ctk.CTkToplevel()
        self.splash.overrideredirect(True) # Quitar bordes y botones
        self.splash.attributes("-topmost", True)
        self.splash.attributes("-alpha", 0.0) # Empezar invisible
        self.splash.configure(fg_color="black")

        # Cargar Imagen Maxtrong
        ruta_img = os.path.join(os.path.dirname(__file__), "maxtrong pacs.png")
        try:
            img_maxtrong = Image.open(ruta_img)
            # Ajustar tamaño (ejemplo 800x450 o similar)
            img_maxtrong = img_maxtrong.resize((900, 500), Image.LANCZOS)
            self.img_splash = ImageTk.PhotoImage(img_maxtrong)
            lbl = ctk.CTkLabel(self.splash, image=self.img_splash, text="")
            lbl.pack()
        except:
            # Si no encuentra la imagen, poner un texto elegante
            ctk.CTkLabel(self.splash, text="MAXTRONG PACS\n2026", font=("Segoe UI", 50, "bold"), text_color="#17a2b8").pack(pady=100, padx=100)

        # Centrar Splash
        sw = self.splash.winfo_screenwidth()
        sh = self.splash.winfo_screenheight()
        x = (sw // 2) - 450
        y = (sh // 2) - 250
        self.splash.geometry(f"900x500+{x}+{y}")

        # --- ANIMACIÓN ---
        self.fade_in()

    def fade_in(self):
        alpha = self.splash.attributes("-alpha")
        if alpha < 1.0:
            alpha += 0.05
            self.splash.attributes("-alpha", alpha)
            self.after(40, self.fade_in)
        else:
            # Mantener 3 segundos materializado
            self.after(3000, self.fade_out)

    def fade_out(self):
        alpha = self.splash.attributes("-alpha")
        if alpha > 0.0:
            alpha -= 0.05
            self.splash.attributes("-alpha", alpha)
            self.after(40, self.fade_out)
        else:
            self.splash.destroy()
            self.inicializar_pacs()

    def inicializar_pacs(self):
        # Restaurar configuración de la ventana principal
        self.title("SISTEMA RADIOLOGICO PROFESIONAL V7.5 - MAXTRONG")
        self.geometry("1500x950")
        self.deiconify() # Mostrar ventana principal

        # --- LÓGICA ORIGINAL REPARADA ---
        self.ruta_storage = os.path.join(os.path.expanduser("~"), "Desktop", "SISTEMA_PACS_DATOS")
        os.makedirs(os.path.join(self.ruta_storage, "INFORMADOS"), exist_ok=True)
        self.db_path = os.path.join(self.ruta_storage, "database.json")
        self.cargar_db()
        
        self.paciente_actual = None
        self.lista_rutas_img = []
        self.indice_img = 0
        self.img_base = None
        self.img_final = None
        self.ds_actual = None
        self.modo_lupa = False
        self.modo_regla = False
        self.medicion_puntos = []

        # --- INTERFAZ ---
        self.configurar_interfaz()

    def configurar_interfaz(self):
        # SIDEBAR IZQUIERDA
        self.sidebar_l = ctk.CTkFrame(self, width=300, corner_radius=CORNER_RADIUS)
        self.sidebar_l.pack(side="left", fill="y", padx=15, pady=15)
        
        ctk.CTkButton(self.sidebar_l, text="+ CARGAR ESTUDIO", font=FONT_BOLD, fg_color="#28a745", corner_radius=CORNER_RADIUS, command=self.importar).pack(pady=15, padx=15, fill="x")
        
        self.tab_pacs = ctk.CTkTabview(self.sidebar_l, corner_radius=CORNER_RADIUS)
        self.tab_pacs.pack(fill="both", expand=True, padx=5, pady=5)
        self.tab_pacs.add("PENDIENTES")
        self.tab_pacs.add("INFORMADOS")
        
        self.f_pendientes = ctk.CTkScrollableFrame(self.tab_pacs.tab("PENDIENTES"), fg_color="transparent")
        self.f_pendientes.pack(fill="both", expand=True)
        self.f_informados = ctk.CTkScrollableFrame(self.tab_pacs.tab("INFORMADOS"), fg_color="transparent")
        self.f_informados.pack(fill="both", expand=True)
        self.actualizar_listas_ui()

        # SIDEBAR DERECHA
        self.sidebar_r = ctk.CTkFrame(self, width=220, corner_radius=CORNER_RADIUS)
        self.sidebar_r.pack(side="right", fill="y", padx=15, pady=15)

        ctk.CTkLabel(self.sidebar_r, text="EXPLORACIÓN", font=FONT_BOLD).pack(pady=15)
        tool_frame = ctk.CTkFrame(self.sidebar_r, fg_color="transparent")
        tool_frame.pack(pady=10)

        self.btn_lupa = ctk.CTkButton(tool_frame, text="🔍", font=FONT_ICONS, width=75, height=75, corner_radius=CORNER_RADIUS, fg_color="#6c757d", command=self.toggle_lupa)
        self.btn_lupa.grid(row=0, column=0, padx=10)
        
        self.btn_regla = ctk.CTkButton(tool_frame, text="📏", font=FONT_ICONS, width=75, height=75, corner_radius=CORNER_RADIUS, fg_color="#6c757d", command=self.toggle_regla)
        self.btn_regla.grid(row=0, column=1, padx=10)

        ctk.CTkLabel(self.sidebar_r, text="AJUSTES", font=FONT_BOLD).pack(pady=25)
        ctk.CTkLabel(self.sidebar_r, text="☀️", font=FONT_ICONS).pack()
        self.slide_brillo = ctk.CTkSlider(self.sidebar_r, from_=-100, to=100, command=self.procesar_ver)
        self.slide_brillo.set(0); self.slide_brillo.pack(padx=15, pady=5)
        
        ctk.CTkLabel(self.sidebar_r, text="◑", font=FONT_ICONS).pack(pady=(15,0))
        self.slide_contra = ctk.CTkSlider(self.sidebar_r, from_=0.5, to=3.0, command=self.procesar_ver)
        self.slide_contra.set(1.0); self.slide_contra.pack(padx=15, pady=5)

        ctk.CTkButton(self.sidebar_r, text="REDACTAR INFORME", font=FONT_BOLD, fg_color="#17a2b8", corner_radius=CORNER_RADIUS, command=self.ventana_informe).pack(side="bottom", pady=20, padx=15, fill="x")
        ctk.CTkButton(self.sidebar_r, text="ELIMINAR ESTUDIO", font=FONT_NORMAL, fg_color="#dc3545", corner_radius=CORNER_RADIUS, command=self.eliminar_estudio).pack(side="bottom", pady=5, padx=15, fill="x")

        # VISOR
        self.container = ctk.CTkFrame(self, fg_color="black", corner_radius=CORNER_RADIUS)
        self.container.pack(side="left", fill="both", expand=True, padx=5, pady=15)
        self.visor = ctk.CTkLabel(self.container, text="")
        self.visor.pack(fill="both", expand=True, padx=5, pady=5)
        self.visor.bind("<Motion>", self.mover_lupa)
        self.visor.bind("<Button-1>", self.click_medicion)

        self.nav = ctk.CTkFrame(self.container, height=45, fg_color="#343a40", corner_radius=CORNER_RADIUS)
        self.nav.pack(side="bottom", fill="x", padx=10, pady=10)
        ctk.CTkButton(self.nav, text="❮❮", width=50, fg_color="transparent", command=self.atras).pack(side="left", padx=15)
        self.lbl_idx = ctk.CTkLabel(self.nav, text="0/0", font=FONT_BOLD); self.lbl_idx.pack(side="left", expand=True)
        ctk.CTkButton(self.nav, text="❯❯", width=50, fg_color="transparent", command=self.adelante).pack(side="right", padx=15)

    # --- LÓGICA INTERNA ---
    def cargar_db(self):
        if os.path.exists(self.db_path):
            with open(self.db_path, "r") as f: self.db = json.load(f)
        else: self.db = {"PENDIENTES": {}, "INFORMADOS": {}}

    def guardar_db(self):
        with open(self.db_path, "w") as f: json.dump(self.db, f)

    def actualizar_listas_ui(self):
        for w in self.f_pendientes.winfo_children(): w.destroy()
        for w in self.f_informados.winfo_children(): w.destroy()
        for p in self.db["PENDIENTES"]:
            ctk.CTkButton(self.f_pendientes, text=p, anchor="w", fg_color="transparent", command=lambda n=p: self.cargar_estudio(n, "PENDIENTES")).pack(fill="x")
        for p in self.db["INFORMADOS"]:
            ctk.CTkButton(self.f_informados, text=p, anchor="w", fg_color="transparent", text_color="#aaa", command=lambda n=p: self.cargar_estudio(n, "INFORMADOS")).pack(fill="x")

    def importar(self):
        folder = filedialog.askdirectory()
        if not folder: return
        nombre = os.path.basename(folder)
        rutas = [os.path.join(folder, f) for f in os.listdir(folder) if f.lower().endswith(('.dcm', '.jpg', '.png'))]
        if rutas:
            self.db["PENDIENTES"][nombre] = rutas
            self.guardar_db(); self.actualizar_listas_ui()

    def cargar_estudio(self, nombre, categoria):
        self.paciente_actual = nombre; self.lista_rutas_img = self.db[categoria][nombre]
        self.indice_img = 0; self.cargar_img()

    def cargar_img(self):
        ruta = self.lista_rutas_img[self.indice_img]
        if ruta.lower().endswith('.dcm'):
            self.ds_actual = pydicom.dcmread(ruta); img = self.ds_actual.pixel_array.astype(float)
            self.img_base = np.uint8((np.maximum(img,0) / img.max()) * 255.0)
        else: self.img_base = cv2.imread(ruta, cv2.IMREAD_GRAYSCALE); self.ds_actual = None
        self.procesar_ver()

    def procesar_ver(self, *args):
        if self.img_base is None: return
        img = cv2.convertScaleAbs(self.img_base, alpha=self.slide_contra.get(), beta=self.slide_brillo.get())
        self.img_final = cv2.bitwise_not(img) # Siempre Negativo
        self.dibujar()

    def dibujar(self, lupa_pos=None):
        if self.img_final is None: return
        h_orig, w_orig = self.img_final.shape
        img_pil = Image.fromarray(self.img_final).convert("RGB")
        vw, vh = self.visor.winfo_width(), self.visor.winfo_height()
        if vw < 10: vw, vh = 1000, 800
        ratio = min(vw/w_orig, vh/h_orig)
        new_w, new_h = int(w_orig * ratio), int(h_orig * ratio)
        img_v = img_pil.resize((new_w, new_h), Image.LANCZOS)
        bg = Image.new("RGB", (vw, vh), (0, 0, 0))
        off_x, off_y = (vw - new_w)//2, (vh - new_h)//2
        bg.paste(img_v, (off_x, off_y))

        if self.modo_lupa and lupa_pos:
            x, y = lupa_pos; zoom = 2; tam = 180
            rx, ry = (x-off_x)/ratio, (y-off_y)/ratio
            if 0 <= rx <= w_orig and 0 <= ry <= h_orig:
                crop = img_pil.crop((rx-tam/zoom/ratio, ry-tam/zoom/ratio, rx+tam/zoom/ratio, ry+tam/zoom/ratio))
                bg.paste(crop.resize((tam*2, tam*2), Image.LANCZOS), (x-tam, y-tam))

        if self.modo_regla and len(self.medicion_puntos) == 2:
            draw = ImageDraw.Draw(bg)
            p1, p2 = self.medicion_puntos
            draw.line([p1, p2], fill="#00ffff", width=3)
            dist_px_real = (np.sqrt((p1[0]-p2[0])**2 + (p1[1]-p2[1])**2)) / ratio
            spacing = float(self.ds_actual.PixelSpacing[0]) if (self.ds_actual and hasattr(self.ds_actual, 'PixelSpacing')) else 0.264
            dist_mm = dist_px_real * spacing
            draw.text((p2[0], p2[1]+15), f"{dist_mm:.2f} mm", fill="yellow")

        self.tk_img = ImageTk.PhotoImage(bg)
        self.visor.configure(image=self.tk_img)
        self.lbl_idx.configure(text=f"{self.indice_img+1}/{len(self.lista_rutas_img)}")

    def toggle_lupa(self):
        self.modo_lupa = not self.modo_lupa; self.modo_regla = False
        self.btn_lupa.configure(fg_color="#17a2b8" if self.modo_lupa else "#6c757d")
        self.btn_regla.configure(fg_color="#6c757d"); self.dibujar()

    def toggle_regla(self):
        self.modo_regla = not self.modo_regla; self.modo_lupa = False; self.medicion_puntos = []
        self.btn_regla.configure(fg_color="#17a2b8" if self.modo_regla else "#6c757d")
        self.btn_lupa.configure(fg_color="#6c757d"); self.dibujar()

    def mover_lupa(self, e): 
        if self.modo_lupa: self.dibujar(lupa_pos=(e.x, e.y))
    def click_medicion(self, e):
        if self.modo_regla:
            self.medicion_puntos.append((e.x, e.y))
            if len(self.medicion_puntos) > 2: self.medicion_puntos = [self.medicion_puntos[-1]]
            self.dibujar()
    def adelante(self): self.indice_img = min(len(self.lista_rutas_img)-1, self.indice_img+1); self.cargar_img()
    def atras(self): self.indice_img = max(0, self.indice_img-1); self.cargar_img()

    def ventana_informe(self):
        if not self.paciente_actual: return
        from docx import Document; from docx.shared import Pt
        win = ctk.CTkToplevel(self); win.title("INFORME"); win.geometry("1000x900"); win.attributes("-topmost", True)
        side_b = ctk.CTkFrame(win, width=170, corner_radius=CORNER_RADIUS, fg_color="#343a40")
        side_b.pack(side="right", fill="y", padx=15, pady=15)
        txt = scrolledtext.ScrolledText(win, font=("Arial Narrow", 14), wrap="word", bg="#1e1e1e", fg="white")
        txt.pack(side="left", fill="both", expand=True, padx=20, pady=20)
        datos = {'r': getattr(self.ds_actual, 'PatientID', '---') if self.ds_actual else '---', 'ex': getattr(self.ds_actual, 'StudyDescription', 'ESTUDIO') if self.ds_actual else 'ESTUDIO'}
        txt.insert("1.0", f"\n\n\nNOMBRE: {self.paciente_actual}\tEDAD: 53\nFECHA: {datetime.now().strftime('%d/%m/%Y')}\tRUT: {datos['r']}\n\n({datos['ex']})\n\nINFORME:\n\n\n\n(FIRMA)")
        def finalizar():
            doc = Document(); doc.styles['Normal'].font.name = 'Arial Narrow'; doc.styles['Normal'].font.size = Pt(12); doc.add_paragraph("\n" * 3)
            p = doc.add_paragraph(); p.add_run("NOMBRE: ").bold = True; p.add_run(f"{self.paciente_actual}\t\t"); p.add_run("EDAD: ").bold = True; p.add_run("53")
            doc.add_paragraph(txt.get("7.0", "end"))
            p_path = os.path.join(self.ruta_storage, "INFORMADOS", self.paciente_actual.replace(" ","_"))
            os.makedirs(p_path, exist_ok=True); doc.save(os.path.join(p_path, f"INFORME_{self.paciente_actual}.docx"))
            if self.paciente_actual in self.db["PENDIENTES"]:
                self.db["INFORMADOS"][self.paciente_actual] = self.db["PENDIENTES"].pop(self.paciente_actual)
                self.guardar_db(); self.actualizar_listas_ui()
            win.destroy(); messagebox.showinfo("EXITO", "Informe guardado.")
        ctk.CTkButton(side_b, text="💾\nGUARDAR", font=FONT_BOLD, fg_color="#28a745", height=100, command=finalizar).pack(pady=25, padx=15, fill="x")

    def eliminar_estudio(self):
        if not self.paciente_actual: return
        if messagebox.askyesno("ELIMINAR", f"¿Eliminar estudio de {self.paciente_actual}?"):
            if self.paciente_actual in self.db["PENDIENTES"]: self.db["PENDIENTES"].pop(self.paciente_actual)
            elif self.paciente_actual in self.db["INFORMADOS"]: 
                self.db["INFORMADOS"].pop(self.paciente_actual)
                shutil.rmtree(os.path.join(self.ruta_storage, "INFORMADOS", self.paciente_actual.replace(" ","_")), ignore_errors=True)
            self.guardar_db(); self.actualizar_listas_ui(); self.img_base = None; self.visor.configure(image=""); self.paciente_actual = None

if __name__ == "__main__":
    app = EstacionRadiologica(); app.mainloop()